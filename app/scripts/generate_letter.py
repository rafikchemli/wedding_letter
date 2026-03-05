"""
Generate IRCC invitation letter DOCX using python-docx.
Reads JSON from stdin, writes DOCX to stdout.
Usage: echo '{"fullName":"...", ...}' | python generate_letter.py [fr|en]
"""
import sys
import json
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Twips, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MONTHS_FR = ['janvier','février','mars','avril','mai','juin','juillet',
             'août','septembre','octobre','novembre','décembre']
MONTHS_EN = ['January','February','March','April','May','June','July',
             'August','September','October','November','December']

def fmt_fr(iso):
    d = datetime.fromisoformat(iso)
    return f"{d.day} {MONTHS_FR[d.month-1]} {d.year}"

def fmt_en(iso):
    d = datetime.fromisoformat(iso)
    return f"{MONTHS_EN[d.month-1]} {d.day}, {d.year}"

def add_run(para, text, bold=False, size=Pt(11)):
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = size
    r.bold = bold
    # Force Times New Roman on East Asian fallback too
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return r

def set_spacing(para, after=Twips(80), before=Twips(0)):
    pf = para.paragraph_format
    pf.space_after = after
    pf.space_before = before
    pf.line_spacing = 1.0  # single spacing (lineRule="auto")

def make_letter(data, lang='fr'):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Twips(1080)
    section.bottom_margin = Twips(720)
    section.left_margin = Twips(1200)
    section.right_margin = Twips(1200)

    # Reset default Normal style (python-docx template has space_after=200)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = None  # inherit (black)
    style.paragraph_format.space_after = Twips(0)
    style.paragraph_format.space_before = Twips(0)
    style.paragraph_format.line_spacing = 1.0

    d = data
    rel_parts = (d.get('relationship') or '').split('|')
    rel_fr = rel_parts[0] if rel_parts else ''
    rel_en = rel_parts[1] if len(rel_parts) > 1 else rel_fr

    is_female = d.get('gender') == 'F'
    title_fr = 'Mme' if is_female else 'M.'
    title_en = 'Mrs.' if is_female else 'Mr.'
    ne = 'née' if is_female else 'né'
    il_elle = 'Elle' if is_female else 'Il'
    he_she = 'She' if is_female else 'He'
    her_his = 'her' if is_female else 'his'
    domiciliee = 'domiciliée' if is_female else 'domicilié'

    # Last name for subsequent references
    name_parts = (d.get('fullName') or '').strip().split()
    last_name = name_parts[-1] if name_parts else d.get('fullName', '')

    passport_fr = f", titulaire du passeport n\u00b0\u00a0{d['passportNumber']} ({d.get('issuingCountry','—')})" if d.get('passportNumber') else ''
    passport_en = f", holder of passport no.\u00a0{d['passportNumber']} ({d.get('issuingCountry','—')})" if d.get('passportNumber') else ''

    today_fr = fmt_fr(datetime.now().strftime('%Y-%m-%d'))
    today_en = fmt_en(datetime.now().strftime('%Y-%m-%d'))

    dob_fr = fmt_fr(d['dob'])
    dob_en = fmt_en(d['dob'])
    arr_fr = fmt_fr(d['arrivalDate'])
    arr_en = fmt_en(d['arrivalDate'])
    dep_fr = fmt_fr(d['departureDate'])
    dep_en = fmt_en(d['departureDate'])

    phone = d.get('phone', '')
    accom_addr = (d.get('accommodationAddress') or '').strip()
    at_host_home = not accom_addr or '308-267' in accom_addr or 'chemli' in accom_addr.lower()

    TWO_LINES = Twips(480)
    TIGHT = Twips(0)
    BODY = Twips(80)

    if lang == 'fr':
        # ── Header ──
        p = doc.add_paragraph(); add_run(p, 'MADJDI RAFIK CHEMLI', bold=True, size=Pt(13)); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, '308-267 Rachel Est'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Montréal (QC), Canada, H2W 1E5'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Tél. : 514-793-1185'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Courriel : rafik.madjdi.chemli@gmail.com'); set_spacing(p, after=Twips(200))
        p = doc.add_paragraph(); add_run(p, f'Montréal, le {today_fr}'); set_spacing(p, after=TWO_LINES)

        # Addressee
        p = doc.add_paragraph(); add_run(p, 'À l\u2019attention de l\u2019agent des visas,'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Immigration, Réfugiés et Citoyenneté Canada (IRCC)'); set_spacing(p, after=Twips(200))

        # Subject
        p = doc.add_paragraph()
        add_run(p, 'Objet\u00a0: ', bold=True)
        add_run(p, f"Lettre d\u2019invitation \u2014 Demande de visa de résident temporaire (visiteur) de {title_fr}\u00a0{d['fullName']}")
        set_spacing(p, after=TWO_LINES)

        # §1 — Host intro
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Je soussigné, '); add_run(p, 'Madjdi Rafik Chemli', bold=True)
        add_run(p, ', citoyen canadien depuis 2006, résidant au Canada de façon continue depuis 2014, actuellement employé à titre de ')
        add_run(p, 'Senior AI Engineer', bold=True); add_run(p, ' chez ')
        add_run(p, 'NewMathData', bold=True)
        add_run(p, ', demeurant à l\u2019adresse indiquée ci\u2011dessus, invite par la présente\u00a0:')
        set_spacing(p, after=BODY)

        # §2 — Visitor intro
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'{title_fr}\u00a0'); add_run(p, d['fullName'], bold=True)
        add_run(p, f', {ne} le '); add_run(p, dob_fr, bold=True)
        add_run(p, f'{passport_fr}, {domiciliee} au '); add_run(p, d['address'], bold=True)
        add_run(p, f', téléphone\u00a0: {phone}.')
        set_spacing(p, after=BODY)

        # §3 — Relationship
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'Je précise que {title_fr}\u00a0{last_name} est {rel_fr} et que cette invitation est faite dans le seul but d\u2019une visite temporaire au Canada.')
        set_spacing(p, after=BODY)

        # §4 — Visit purpose + dates
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'La présente visite a pour objet de permettre à {title_fr}\u00a0{last_name} d\u2019assister à mon mariage, prévu le ')
        add_run(p, '19 septembre 2026', bold=True)
        add_run(p, ' à la salle '); add_run(p, 'L\u2019Éloi, à Montréal (Québec)', bold=True)
        add_run(p, '. La période de séjour prévue est du '); add_run(p, arr_fr, bold=True)
        add_run(p, ' au '); add_run(p, dep_fr, bold=True)
        add_run(p, f" ({d['duration']}). {il_elle} quittera le Canada au plus tard le {dep_fr}, afin de retourner en {d['returnCountry']}.")
        set_spacing(p, after=BODY)

        # §5 — Accommodation + financial
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if at_host_home:
            add_run(p, f'Pendant la durée de son séjour, {title_fr}\u00a0{last_name} logera à mon domicile, à l\u2019adresse indiquée en en\u2011tête. Je confirme que je prendrai en charge l\u2019hébergement pour la totalité du séjour.')
        else:
            add_run(p, f'Pendant la durée de son séjour, {title_fr}\u00a0{last_name} logera à l\u2019adresse suivante\u00a0: {accom_addr}.')
        add_run(p, f' {title_fr}\u00a0{last_name} assumera toutes les autres dépenses liées au voyage et au séjour, notamment\u00a0: billets d\u2019avion (aller\u2011retour), transport, repas et dépenses personnelles.')
        set_spacing(p, after=BODY)

        # §6 — Ties
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'À ma connaissance, {title_fr}\u00a0{last_name} dispose d\u2019attaches solides dans son pays de résidence, notamment\u00a0: ')
        add_run(p, d['returnReason'], bold=True)
        add_run(p, f', et {il_elle.lower()} a l\u2019intention de reprendre ses obligations professionnelles à l\u2019issue de son séjour autorisé.')
        set_spacing(p, after=BODY)

        # §7 — Compliance
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'Je comprends que ce séjour est strictement temporaire. {title_fr}\u00a0{last_name} s\u2019engage à respecter les conditions applicables aux visiteurs, notamment à ne pas travailler et à ne pas étudier au Canada sans autorisation.')
        set_spacing(p, after=BODY)

        # §8 — Contact + Thanks
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Je demeure à votre disposition pour tout renseignement complémentaire. Merci de l\u2019attention portée à la présente.')
        set_spacing(p, after=TWO_LINES)

        # Closing
        p = doc.add_paragraph()
        add_run(p, 'Veuillez agréer, Madame, Monsieur, l\u2019expression de mes salutations distinguées.')
        set_spacing(p, after=Twips(200))

        # Signature
        p = doc.add_paragraph(); add_run(p, 'Madjdi Rafik Chemli', bold=True); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Senior AI Engineer — NewMathData'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Citoyen canadien'); set_spacing(p, after=TWO_LINES)

        # Enclosures
        p = doc.add_paragraph()
        r = add_run(p, 'Pièces jointes (de l\u2019hôte) — copies\u00a0:', bold=True); r.underline = True
        set_spacing(p, after=Twips(40))
        for item in [
            'Copie preuve de citoyenneté\u00a0: passeport canadien (page d\u2019identité) et/ou certificat',
            'Preuve d\u2019adresse au Canada\u00a0: facture de services publics',
            'Lettre d\u2019emploi / preuves d\u2019emploi',
            'Preuve liée au mariage\u00a0: réservation de salle / confirmation',
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Twips(360)
            add_run(p, item)
            set_spacing(p, after=Twips(20))

    else:  # English
        # ── Header ──
        p = doc.add_paragraph(); add_run(p, 'MADJDI RAFIK CHEMLI', bold=True, size=Pt(13)); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, '308-267 Rachel Est'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Montréal (QC), Canada, H2W 1E5'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Phone: 514-793-1185'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Email: rafik.madjdi.chemli@gmail.com'); set_spacing(p, after=Twips(200))
        p = doc.add_paragraph(); add_run(p, f'Montréal, {today_en}'); set_spacing(p, after=TWO_LINES)

        # Addressee
        p = doc.add_paragraph(); add_run(p, 'To: Visa Officer'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Immigration, Refugees and Citizenship Canada (IRCC)'); set_spacing(p, after=Twips(200))

        # Subject
        p = doc.add_paragraph()
        add_run(p, 'Subject: ', bold=True)
        add_run(p, f"Letter of Invitation \u2014 Temporary Resident Visa (Visitor) application for {title_en} {d['fullName']}")
        set_spacing(p, after=TWO_LINES)

        # §1 — Host intro
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'I, '); add_run(p, 'Madjdi Rafik Chemli', bold=True)
        add_run(p, ', a Canadian citizen since 2006, residing in Canada continuously since 2014, currently employed as a ')
        add_run(p, 'Senior AI Engineer', bold=True); add_run(p, ' with ')
        add_run(p, 'NewMathData', bold=True)
        add_run(p, ', living at the address indicated above, hereby invite:')
        set_spacing(p, after=BODY)

        # §2 — Visitor intro
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'{title_en} '); add_run(p, d['fullName'], bold=True)
        add_run(p, ', born '); add_run(p, dob_en, bold=True)
        add_run(p, f'{passport_en}, residing at '); add_run(p, d['address'], bold=True)
        add_run(p, f', phone: {phone}.')
        set_spacing(p, after=BODY)

        # §3 — Relationship
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'{title_en} {last_name} is {rel_en}, and this invitation is issued solely for a temporary visit to Canada.')
        set_spacing(p, after=BODY)

        # §4 — Visit purpose + dates
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'The purpose of this visit is for {title_en} {last_name} to attend my wedding, scheduled for ')
        add_run(p, 'September 19, 2026', bold=True)
        add_run(p, ' at '); add_run(p, "L\u2019Éloi in Montréal, Quebec", bold=True)
        add_run(p, '. The planned stay is from '); add_run(p, arr_en, bold=True)
        add_run(p, ' to '); add_run(p, dep_en, bold=True)
        add_run(p, f" ({d['duration']}). {he_she} will depart Canada no later than {dep_en} to return to {d['returnCountry']}.")
        set_spacing(p, after=BODY)

        # §5 — Accommodation + financial
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if at_host_home:
            add_run(p, f'During {her_his} stay, {title_en} {last_name} will stay at my home at the address shown in the header of this letter. I confirm I will provide accommodation for the full duration of the visit.')
        else:
            add_run(p, f'During {her_his} stay, {title_en} {last_name} will stay at the following address: {accom_addr}.')
        add_run(p, f' {title_en} {last_name} will cover all other travel and living expenses, including round-trip airfare, transportation, meals, and personal expenses.')
        set_spacing(p, after=BODY)

        # §6 — Ties
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'To the best of my knowledge, {title_en} {last_name} has strong ties to {her_his} country of residence, including: ')
        add_run(p, d['returnReason'], bold=True)
        add_run(p, f', and {he_she.lower()} intends to resume {her_his} professional obligations upon the end of {her_his} authorized stay.')
        set_spacing(p, after=BODY)

        # §7 — Compliance
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, f'I understand this visit is strictly temporary. {title_en} {last_name} will comply with visitor conditions, including not working and not studying in Canada without authorization.')
        set_spacing(p, after=BODY)

        # §8 — Contact + Thanks
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, 'Please feel free to contact me should you require any further information. Thank you for your consideration.')
        set_spacing(p, after=TWO_LINES)

        # Closing
        p = doc.add_paragraph(); add_run(p, 'Sincerely,'); set_spacing(p, after=Twips(200))

        # Signature
        p = doc.add_paragraph(); add_run(p, 'Madjdi Rafik Chemli', bold=True); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Senior AI Engineer — NewMathData'); set_spacing(p, after=TIGHT)
        p = doc.add_paragraph(); add_run(p, 'Canadian citizen'); set_spacing(p, after=TWO_LINES)

        # Enclosures
        p = doc.add_paragraph()
        r = add_run(p, 'Attachments (host) \u2014 copies:', bold=True); r.underline = True
        set_spacing(p, after=Twips(40))
        for item in [
            'Proof of Canadian status: Canadian passport biodata page and/or citizenship certificate',
            'Proof of Canadian address: utility bill / lease / mortgage statement',
            'Employment confirmation / proof of employment',
            'Wedding-related proof: venue booking / confirmation / invitation',
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Twips(360)
            add_run(p, item)
            set_spacing(p, after=Twips(20))

    return doc


if __name__ == '__main__':
    lang = sys.argv[1] if len(sys.argv) > 1 else 'fr'
    data = json.loads(sys.stdin.read())
    doc = make_letter(data, lang)
    buf = BytesIO()
    doc.save(buf)
    sys.stdout.buffer.write(buf.getvalue())
